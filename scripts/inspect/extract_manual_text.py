from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def sanitize_name(value: str) -> str:
    safe = []
    for ch in value:
        if ch.isalnum() or ch in ("-", "_", "."):
            safe.append(ch)
        else:
            safe.append("_")
    return "".join(safe)


def extract_pdf(input_path: Path, out_dir: Path) -> dict:
    reader = PdfReader(str(input_path))
    pages = []
    full_text_parts = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(
            {
                "page_number": index,
                "char_count": len(text),
                "text": text,
            }
        )
        full_text_parts.append(text)

    full_text = "\n\n".join(full_text_parts)
    page_json_path = out_dir / f"{sanitize_name(input_path.stem)}.pages.json"
    text_path = out_dir / f"{sanitize_name(input_path.stem)}.txt"
    page_json_path.write_text(json.dumps(pages, ensure_ascii=False, indent=2), encoding="utf-8")
    text_path.write_text(full_text, encoding="utf-8")
    return {
        "source": str(input_path),
        "type": "pdf",
        "page_count": len(reader.pages),
        "total_chars": len(full_text),
        "text_path": str(text_path),
        "page_json_path": str(page_json_path),
    }


def extract_docx(input_path: Path, out_dir: Path) -> dict:
    document = Document(str(input_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    full_text = "\n".join(paragraphs)
    paragraph_json_path = out_dir / f"{sanitize_name(input_path.stem)}.paragraphs.json"
    text_path = out_dir / f"{sanitize_name(input_path.stem)}.txt"
    paragraph_json_path.write_text(
        json.dumps(
            [
                {
                    "paragraph_number": index + 1,
                    "char_count": len(text),
                    "text": text,
                }
                for index, text in enumerate(paragraphs)
            ],
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    text_path.write_text(full_text, encoding="utf-8")
    return {
        "source": str(input_path),
        "type": "docx",
        "paragraph_count": len(paragraphs),
        "total_chars": len(full_text),
        "text_path": str(text_path),
        "paragraph_json_path": str(paragraph_json_path),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--inputs", nargs="+", required=True)
    parser.add_argument("--out-dir", required=True)
    args = parser.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    manifest = []
    for raw_input in args.inputs:
        input_path = Path(raw_input)
        suffix = input_path.suffix.lower()
        if suffix == ".pdf":
            manifest.append(extract_pdf(input_path, out_dir))
        elif suffix == ".docx":
            manifest.append(extract_docx(input_path, out_dir))
        else:
            manifest.append(
                {
                    "source": str(input_path),
                    "type": "unsupported",
                }
            )

    manifest_path = out_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(manifest_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
